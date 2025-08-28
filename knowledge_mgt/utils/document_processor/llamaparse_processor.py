#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import logging
from typing import List, Dict, Any
import tempfile
from urllib.parse import urlparse
from .base_processor import BaseDocumentProcessor

logger = logging.getLogger(__name__)

class LlamaParseProcessor(BaseDocumentProcessor):
    """使用LlamaParse的文档处理器"""
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)
        self.parser = None
        self._initialize_parser()
    
    def _initialize_parser(self):
        """初始化LlamaParse解析器"""
        try:
            from llama_parse import LlamaParse
            import nest_asyncio
            
            # 设置异步支持
            nest_asyncio.apply()
            
            # 检查是否有 API key
            api_key = os.environ.get('LLAMA_CLOUD_API_KEY') or os.environ.get('LLAMA_PARSE_API_KEY')
            
            if not api_key:
                logger.warning("LlamaParse API key 未设置，将使用降级模式")
                logger.info("如需使用 LlamaParse，请设置环境变量 LLAMA_CLOUD_API_KEY 或 LLAMA_PARSE_API_KEY")
                self.parser = None
                return
            
            # 初始化解析器，支持多种结果类型
            self.parser = LlamaParse(
                api_key=api_key,
                result_type="text"  # "markdown" 和 "text" 可选
            )
            
            # 设置文件提取器，支持多种格式
            self.file_extractor = {
                ".pdf": self.parser,
                ".docx": self.parser,
                ".doc": self.parser,
                ".xlsx": self.parser,  # 添加 Excel 支持
                ".xls": self.parser,   # 添加 Excel 支持
                ".csv": self.parser,   # 添加 CSV 支持
                ".md": self.parser,
                ".txt": self.parser,
                ".html": self.parser,
                ".rtf": self.parser,
                ".odt": self.parser,
                ".pages": self.parser,
                ".epub": self.parser,
                ".mobi": self.parser
            }
            
            logger.info("LlamaParse解析器初始化成功，支持多种文件格式")
            
        except ImportError as e:
            logger.error(f"LlamaParse导入失败: {str(e)}")
            self.parser = None
        except Exception as e:
            logger.error(f"LlamaParse初始化失败: {str(e)}")
            self.parser = None
    
    def extract_text(self, file_path: str) -> str:
        """使用LlamaParse提取文档文本"""
        if not self.parser:
            logger.warning("LlamaParse解析器未初始化，尝试使用降级提取")
            # 若是URL，直接走URL降级；否则走文件降级
            try:
                if str(file_path).lower().startswith(("http://", "https://")):
                    return self._fallback_extract_from_url(str(file_path))
                return self._fallback_extract(file_path)
            except Exception:
                return ""
        
        try:
            logger.info(f"开始使用LlamaParse提取文本: {file_path}")
            # 如果是网址，先下载到临时文件
            is_url = str(file_path).lower().startswith(("http://", "https://"))
            temp_file_path = None
            original_url = file_path if is_url else None  # 保存原始URL
            if is_url:
                try:
                    logger.info(f"检测到URL，开始下载: {file_path}")
                    # 延迟导入，避免硬依赖
                    import requests  # type: ignore
                    response = requests.get(file_path, timeout=30)
                    response.raise_for_status()
                    # 推断扩展名
                    parsed = urlparse(file_path)
                    url_name = os.path.basename(parsed.path)
                    name, ext = os.path.splitext(url_name)
                    if not ext:
                        # 根据Content-Type推断
                        content_type = response.headers.get('Content-Type', '').lower()
                        if 'pdf' in content_type:
                            ext = '.pdf'
                        elif 'msword' in content_type or 'officedocument.wordprocessingml' in content_type:
                            ext = '.pdf'
                        elif 'excel' in content_type or 'officedocument.spreadsheetml' in content_type:
                            ext = '.xlsx'
                        elif 'html' in content_type:
                            ext = '.html'
                        elif 'plain' in content_type or 'text/' in content_type:
                            ext = '.txt'
                        else:
                            ext = '.bin'
                    fd, temp_file_path = tempfile.mkstemp(suffix=ext)
                    with os.fdopen(fd, 'wb') as tmp:
                        tmp.write(response.content)
                    logger.info(f"URL已下载到临时文件: {temp_file_path}")
                    file_path = temp_file_path
                except Exception as e:
                    logger.error(f"下载URL失败，回退到基础提取: {e}")
                    return self._fallback_extract_from_url(str(original_url))
            
            # 获取文件扩展名
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # 检查是否支持该文件格式
            if hasattr(self, 'file_extractor') and self.file_extractor:
                if file_ext not in self.file_extractor:
                    logger.warning(f"LlamaParse不支持的文件格式: {file_ext}")
                    # 尝试降级处理
                    return self._fallback_extract(file_path)
            else:
                logger.warning("LlamaParse文件提取器未初始化，尝试降级处理")
                return self._fallback_extract(file_path)
            
            # 使用LlamaParse解析文档
            documents = self.parser.load_data(file_path)
            
            if not documents:
                logger.warning("LlamaParse未返回任何文档")
                # 如果原始输入是URL，尝试URL降级提取
                if is_url and original_url:
                    return self._fallback_extract_from_url(str(original_url))
                return ""
            
            # 合并所有文档的文本
            text_content = ""
            for doc in documents:
                if hasattr(doc, 'text'):
                    text_content += doc.text + "\n\n"
                elif hasattr(doc, 'content'):
                    text_content += doc.content + "\n\n"
                elif hasattr(doc, 'page_content'):
                    text_content += doc.page_content + "\n\n"
                else:
                    logger.warning(f"文档对象缺少text/content/page_content属性: {type(doc)}")
                    # 尝试获取其他可能的属性
                    for attr in ['body', 'html', 'markdown']:
                        if hasattr(doc, attr):
                            content = getattr(doc, attr)
                            if content:
                                text_content += str(content) + "\n\n"
                                break
            
            logger.info(f"LlamaParse文本提取完成，长度: {len(text_content)}")
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"LlamaParse文本提取失败: {str(e)}")
            # 如果原始输入是URL，优先走URL降级
            try:
                if is_url and original_url:
                    # 传入原始URL字符串
                    return self._fallback_extract_from_url(str(original_url))
            except Exception:
                pass
            # 否则降级到基础文件提取
            return self._fallback_extract(file_path)
        finally:
            # 清理临时文件
            try:
                if 'temp_file_path' in locals() and temp_file_path and os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
            except Exception:
                pass
    
    def _fallback_extract(self, file_path: str) -> str:
        """降级文本提取方法"""
        try:
            logger.info(f"尝试降级文本提取: {file_path}")
            # 如果是URL，先调用URL专用降级
            if str(file_path).lower().startswith(("http://", "https://")):
                return self._fallback_extract_from_url(str(file_path))
            
            # 对于文本文件，直接读取
            if file_path.lower().endswith(('.txt', '.md')):
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            # 对于HTML文件，尝试使用BeautifulSoup
            if file_path.lower().endswith('.html'):
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        html_content = f.read()
                    # 延迟导入，避免硬依赖
                    from bs4 import BeautifulSoup  # type: ignore
                    soup = BeautifulSoup(html_content, 'html.parser')
                    # 去掉script/style
                    for tag in soup(['script', 'style']):
                        tag.decompose()
                    return soup.get_text(separator='\n')
                except ImportError:
                    logger.warning("BeautifulSoup 不可用，无法提取HTML文本")
                    # 尝试直接读取文本
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            return f.read()
                    except Exception:
                        pass
                except Exception as e:
                    logger.error(f"HTML文本提取失败: {str(e)}")
            
            # 对于Excel文件，尝试使用PandasExcelReader
            if file_path.lower().endswith(('.xlsx', '.xls')):
                try:
                    from llama_index.readers.file import PandasExcelReader
                    logger.info("使用PandasExcelReader提取Excel文件内容")
                    
                    # 创建PandasExcelReader实例
                    excel_reader = PandasExcelReader()
                    
                    # 读取Excel文件
                    documents = excel_reader.load_data(file_path)
                    
                    if documents:
                        text_content = ""
                        for doc in documents:
                            if hasattr(doc, 'text'):
                                text_content += doc.text + "\n\n"
                            elif hasattr(doc, 'page_content'):
                                text_content += doc.page_content + "\n\n"
                            elif hasattr(doc, 'content'):
                                text_content += doc.content + "\n\n"
                        
                        logger.info(f"PandasExcelReader提取成功，长度: {len(text_content)}")
                        return text_content.strip()
                    else:
                        logger.warning("PandasExcelReader未返回任何文档")
                        return ""
                        
                except ImportError:
                    try:
                        import pandas as pd
                        logger.info("使用pandas提取Excel文件内容")
                        df = pd.read_excel(file_path, sheet_name=None)  # 读取所有工作表
                        
                        text_content = ""
                        for sheet_name, sheet_df in df.items():
                            text_content += f"\n=== 工作表: {sheet_name} ===\n"
                            text_content += sheet_df.to_string(index=False) + "\n\n"
                        
                        return text_content.strip()
                    except ImportError:
                        try:
                            import openpyxl
                            logger.info("使用openpyxl提取Excel文件内容")
                            workbook = openpyxl.load_workbook(file_path, read_only=True)
                            
                            text_content = ""
                            for sheet_name in workbook.sheetnames:
                                sheet = workbook[sheet_name]
                                text_content += f"\n=== 工作表: {sheet_name} ===\n"
                                
                                for row in sheet.iter_rows(values_only=True):
                                    if any(cell is not None for cell in row):
                                        row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                                        text_content += row_text + "\n"
                                text_content += "\n"
                            
                            workbook.close()
                            return text_content.strip()
                        except ImportError:
                            logger.warning("PandasExcelReader、pandas和openpyxl都不可用，无法提取Excel文本")
                            return ""
                except Exception as e:
                    logger.error(f"PandasExcelReader处理失败: {e}")
                    # 尝试降级到pandas
                    try:
                        import pandas as pd
                        logger.info("降级到pandas处理Excel文件")
                        df = pd.read_excel(file_path, sheet_name=None)
                        
                        text_content = ""
                        for sheet_name, sheet_df in df.items():
                            text_content += f"\n=== 工作表: {sheet_name} ===\n"
                            text_content += sheet_df.to_string(index=False) + "\n\n"
                        
                        return text_content.strip()
                    except Exception as e2:
                        logger.error(f"pandas降级处理也失败: {e2}")
                        return ""
            
            # 对于PDF文件，尝试使用PyPDF2
            if file_path.lower().endswith('.pdf'):
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                        return text
                except ImportError:
                    logger.warning("PyPDF2 不可用，无法提取PDF文本")
            
            # 对于Word文档，尝试使用python-docx
            if file_path.lower().endswith(('.docx', '.doc')):
                try:
                    import docx
                    logger.info("使用python-docx提取Word文档内容")
                    doc = docx.Document(file_path)
                    text_content = ""
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_content += paragraph.text + "\n"
                    return text_content.strip()
                except ImportError:
                    logger.warning("python-docx 不可用，无法提取Word文档文本")
            
            logger.warning(f"降级提取失败，不支持的文件格式: {file_path}")
            return ""
            
        except Exception as e:
            logger.error(f"降级提取也失败: {str(e)}")
            return ""

    def _fallback_extract_from_url(self, url: str) -> str:
        """URL的降级文本提取：下载并基于内容类型做简单解析"""
        try:
            logger.info(f"开始对URL进行降级提取: {url}")
            # 延迟导入，避免硬依赖
            import requests  # type: ignore
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            content_type = response.headers.get('Content-Type', '').lower()
            content = response.content
            # 对常见类型做快速处理
            if 'text/' in content_type or 'json' in content_type or 'xml' in content_type or 'html' in content_type:
                try:
                    text = content.decode('utf-8')
                except Exception:
                    try:
                        text = content.decode('gbk')
                    except Exception:
                        text = content.decode(errors='ignore')
                # 如果是HTML，尽量去除标签
                if 'html' in content_type:
                    try:
                        # 延迟导入，避免硬依赖
                        from bs4 import BeautifulSoup  # type: ignore
                        soup = BeautifulSoup(text, 'html.parser')
                        # 去掉script/style
                        for tag in soup(['script', 'style']):
                            tag.decompose()
                        return soup.get_text(separator='\n')
                    except Exception:
                        return text
                return text
            # 其他二进制类型，写入临时文件后走文件降级
            fd, temp_path = tempfile.mkstemp()
            with os.fdopen(fd, 'wb') as tmp:
                tmp.write(content)
            try:
                return self._fallback_extract(temp_path)
            finally:
                try:
                    os.remove(temp_path)
                except Exception:
                    pass
        except Exception as e:
            logger.error(f"URL降级提取失败: {e}")
            return ""
    
    def split_text(self, text: str) -> List[str]:
        """分割文本为块"""
        if not text:
            return []
        
        try:
            # 使用LlamaParse的智能分割
            try:
                # 尝试新版本的导入方式
                from llama_index.core import Document
                from llama_index.core.node_parser import SentenceSplitter
            except ImportError:
                try:
                    # 尝试旧版本的导入方式
                    from llama_index import Document
                    from llama_index.node_parser import SentenceSplitter
                except ImportError:
                    # 如果都失败，使用兼容性导入
                    try:
                        from llama_index.schema import Document
                        from llama_index.node_parser import SentenceSplitter
                    except ImportError:
                        # 最后的兼容性尝试
                        from llama_index import Document
                        from llama_index.node_parser import SentenceSplitter
            
            # 创建文档对象
            doc = Document(text=text)
            
            # 使用句子分割器
            splitter = SentenceSplitter(
                chunk_size=self.config.get('chunk_size', 1000),
                chunk_overlap=self.config.get('overlap_size', 200)
            )
            
            # 分割文档
            nodes = splitter.get_nodes_from_documents([doc])
            chunks = [node.text for node in nodes]
            
            logger.info(f"LlamaParse文本分割完成，生成 {len(chunks)} 个块")
            return chunks
            
        except Exception as e:
            logger.error(f"LlamaParse文本分割失败: {str(e)}")
            # 降级到基础分割
            return self._fallback_split(text)
    
    def _fallback_split(self, text: str) -> List[str]:
        """降级分割方法"""
        try:
            # 简单的段落分割
            paragraphs = text.split('\n\n')
            chunks = []
            current_chunk = ""
            
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue
                
                if len(current_chunk) + len(para) <= self.config.get('chunk_size', 1000):
                    current_chunk += para + "\n\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = para + "\n\n"
            
            if current_chunk:
                chunks.append(current_chunk.strip())
            
            logger.info(f"降级分割完成，生成 {len(chunks)} 个块")
            return chunks
            
        except Exception as e:
            logger.error(f"降级分割也失败: {str(e)}")
            return [text]  # 返回原文本作为单个块
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文档格式"""
        return [
            # 文档格式
            '.pdf', '.docx', '.doc', '.txt', '.md', '.html', '.rtf',
            '.odt', '.pages', '.epub', '.mobi',
            
            # 电子表格格式
            '.xlsx', '.xls', '.csv',
            
            # 演示文稿格式
            '.pptx', '.ppt',
            
            # 其他格式
            '.xml', '.json', '.yaml', '.yml',
            # 虚拟格式：URL
            'http', 'https'
        ]
    
    def is_available(self) -> bool:
        """检查处理器是否可用"""
        return self.parser is not None
    
    def can_process(self, file_path: str) -> bool:
        """检查是否可以处理指定文件"""
        if not self.is_available():
            return False
        
        # URL直接支持
        if str(file_path).lower().startswith(("http://", "https://")):
            return True
        
        file_ext = os.path.splitext(file_path)[1].lower()
        supported_formats = self.get_supported_formats()
        
        # 检查文件扩展名
        if file_ext in supported_formats:
            return True
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.warning(f"文件不存在: {file_path}")
            return False
        
        # 对于某些特殊格式，可以进一步检查文件头
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)  # 读取前8字节
                
                # PDF文件头检查
                if header.startswith(b'%PDF'):
                    return True
                
                # ZIP文件头检查（docx, xlsx, pptx等）
                if header.startswith(b'PK'):
                    return True
                
                # Office文档文件头检查
                if header.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                    return True
                
        except Exception as e:
            logger.debug(f"文件头检查失败: {str(e)}")
        
        return False
    
    def extract_structure(self, file_path: str) -> Dict[str, Any]:
        """提取文档结构（简化实现）"""
        try:
            # 对于LlamaParse，我们主要关注文本内容
            # 结构信息可以通过其他方式获取
            text_content = self.extract_text(file_path)
            return {
                "text": text_content,
                "length": len(text_content),
                "has_content": bool(text_content.strip())
            }
        except Exception as e:
            logger.error(f"提取文档结构失败: {str(e)}")
            return {"error": str(e)}
