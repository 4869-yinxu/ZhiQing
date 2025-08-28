"""
基于LlamaIndex的多格式文档处理器
支持docx, pdf, doc, md, excel等格式
"""

import os
import logging
from typing import Dict, List, Any, Optional
from pathlib import Path

try:
    # 优先兼容新版本路径 (>=0.10)
    try:
        from llama_index.core.readers import (
            PDFReader,
            DocxReader,
            MarkdownReader,
            CSVReader,
        )
    except Exception:
        # 兼容旧版本路径
        from llama_index.readers.file import (
            PDFReader,
            DocxReader,
            MarkdownReader,
            CSVReader,
        )

    # ExcelReader 位置在不同版本中不同，逐级尝试
    ExcelReader = None
    for path in (
        "llama_index.core.readers.excel",
        "llama_index.readers.excel",
        "llama_index.readers",
    ):
        try:
            module = __import__(path, fromlist=["ExcelReader"])
            ExcelReader = getattr(module, "ExcelReader", None)
            if ExcelReader:
                break
        except Exception:
            continue
    if ExcelReader is None:
        logging.info("LlamaIndex: ExcelReader 未找到，Excel文件处理将被跳过")

    # 新旧版本的 Document 类型兼容（未实际用到时可忽略）
    try:
        from llama_index.core.schema import Document as LlamaDocument  # new
    except Exception:
        try:
            from llama_index.schema import Document as LlamaDocument  # old
        except Exception:
            LlamaDocument = None

    LLAMAINDEX_AVAILABLE = True
except Exception as e:
    LLAMAINDEX_AVAILABLE = False
    logging.info(f"LlamaIndex 导入不可用，已跳过（不影响其他处理器）: {e}")

from .base_processor import BaseDocumentProcessor

logger = logging.getLogger(__name__)


class LlamaIndexProcessor(BaseDocumentProcessor):
    """基于LlamaIndex的文档处理器"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        
        if not LLAMAINDEX_AVAILABLE:
            raise ImportError("LlamaIndex未安装，无法使用此处理器")
        
        # 支持的文件格式
        self.supported_formats = [
            '.txt', '.md', '.markdown',  # 文本格式
            '.pdf',                       # PDF格式
            '.docx', '.doc',             # Word格式
            '.xlsx', '.xls', '.csv'      # Excel格式
        ]
        
        # 格式映射到对应的Reader
        self.reader_mapping = {
            '.txt': self._read_text_file,
            '.md': self._read_markdown_file,
            '.markdown': self._read_markdown_file,
            '.pdf': self._read_pdf_file,
            '.docx': self._read_docx_file,
            '.doc': self._read_doc_file,
            '.xlsx': self._read_excel_file,
            '.xls': self._read_excel_file,
            '.csv': self._read_csv_file
        }
        
        # 处理器配置
        self.chunk_size = config.get('chunk_size', 1000) if config else 1000
        self.chunk_overlap = config.get('chunk_overlap', 200) if config else 200
        
    def is_available(self) -> bool:
        """检查处理器是否可用"""
        return LLAMAINDEX_AVAILABLE
    
    def can_process(self, file_path: str) -> bool:
        """检查是否可以处理指定文件"""
        if not self.validate_file(file_path):
            return False
            
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.supported_formats
    
    def extract_text(self, file_path: str) -> str:
        """从文件中提取文本内容"""
        try:
            self.log_processing(file_path, "开始提取文本")
            
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.reader_mapping:
                raise ValueError(f"不支持的文件格式: {file_ext}")
            
            # 使用对应的Reader提取文本
            reader_func = self.reader_mapping[file_ext]
            text_content = reader_func(file_path)
            
            self.log_processing(file_path, "文本提取完成")
            return text_content
            
        except Exception as e:
            self.handle_error(file_path, e, "文本提取")
    
    def extract_structure(self, file_path: str) -> Dict[str, Any]:
        """提取文档结构信息"""
        try:
            self.log_processing(file_path, "开始提取文档结构")
            
            file_ext = Path(file_path).suffix.lower()
            structure_info = {
                'file_path': file_path,
                'file_type': file_ext,
                'processor': self.processor_name,
                'headings': [],
                'sections': [],
                'metadata': {}
            }
            
            # 根据文件类型提取结构
            if file_ext in ['.docx', '.doc']:
                structure_info.update(self._extract_word_structure(file_path))
            elif file_ext == '.pdf':
                structure_info.update(self._extract_pdf_structure(file_path))
            elif file_ext in ['.xlsx', '.xls']:
                structure_info.update(self._extract_excel_structure(file_path))
            elif file_ext in ['.md', '.markdown']:
                structure_info.update(self._extract_markdown_structure(file_path))
            
            self.log_processing(file_path, "文档结构提取完成")
            return structure_info
            
        except Exception as e:
            self.handle_error(file_path, e, "文档结构提取")
    
    def _read_text_file(self, file_path: str) -> str:
        """读取文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _read_markdown_file(self, file_path: str) -> str:
        """读取Markdown文件"""
        reader = MarkdownReader()
        documents = reader.load_data(file_path)
        return '\n\n'.join([doc.text for doc in documents])
    
    def _read_pdf_file(self, file_path: str) -> str:
        """读取PDF文件"""
        reader = PDFReader()
        documents = reader.load_data(file_path)
        return '\n\n'.join([doc.text for doc in documents])
    
    def _read_docx_file(self, file_path: str) -> str:
        """读取DOCX文件"""
        reader = DocxReader()
        documents = reader.load_data(file_path)
        return '\n\n'.join([doc.text for doc in documents])
    
    def _read_doc_file(self, file_path: str) -> str:
        """读取DOC文件（转换为DOCX后处理）"""
        # 注意：DOC文件需要先转换为DOCX格式
        # 这里可以使用python-docx2txt或其他库
        try:
            import docx2txt
            return docx2txt.process(file_path)
        except ImportError:
            # 如果没有docx2txt，尝试使用其他方法
            logger.warning("docx2txt未安装，DOC文件处理可能受限")
            return self._read_docx_file(file_path)
    
    def _read_excel_file(self, file_path: str) -> str:
        """读取Excel文件"""
        if ExcelReader is None:
            raise ImportError("ExcelReader不可用，无法处理Excel文件")
        reader = ExcelReader()
        documents = reader.load_data(file_path)
        return '\n\n'.join([doc.text for doc in documents])
    
    def _read_csv_file(self, file_path: str) -> str:
        """读取CSV文件"""
        reader = CSVReader()
        documents = reader.load_data(file_path)
        return '\n\n'.join([doc.text for doc in documents])
    
    def _extract_word_structure(self, file_path: str) -> Dict[str, Any]:
        """提取Word文档结构"""
        try:
            import docx
            doc = docx.Document(file_path)
            
            headings = []
            sections = []
            
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                    headings.append({
                        'level': level,
                        'text': para.text,
                        'style': para.style.name
                    })
                elif para.text.strip():
                    sections.append({
                        'type': 'paragraph',
                        'text': para.text
                    })
            
            return {
                'headings': headings,
                'sections': sections,
                'paragraph_count': len(doc.paragraphs)
            }
        except Exception as e:
            logger.warning(f"提取Word文档结构失败: {str(e)}")
            return {'headings': [], 'sections': []}
    
    def _extract_pdf_structure(self, file_path: str) -> Dict[str, Any]:
        """提取PDF文档结构"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            
            headings = []
            sections = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                blocks = page.get_text("dict")["blocks"]
                
                for block in blocks:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text = span["text"].strip()
                                if text:
                                    # 简单的标题检测（基于字体大小）
                                    if span["size"] > 14:  # 假设大于14pt的是标题
                                        headings.append({
                                            'level': 1,
                                            'text': text,
                                            'page': page_num + 1
                                        })
                                    else:
                                        sections.append({
                                            'type': 'text',
                                            'text': text,
                                            'page': page_num + 1
                                        })
            
            doc.close()
            return {
                'headings': headings,
                'sections': sections,
                'page_count': len(doc)
            }
        except Exception as e:
            logger.warning(f"提取PDF文档结构失败: {str(e)}")
            return {'headings': [], 'sections': []}
    
    def _extract_excel_structure(self, file_path: str) -> Dict[str, Any]:
        """提取Excel文档结构"""
        try:
            import pandas as pd
            
            # 读取所有工作表
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            sections = []
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sections.append({
                    'type': 'worksheet',
                    'name': sheet_name,
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist()
                })
            
            return {
                'headings': [],
                'sections': sections,
                'worksheet_count': len(sheet_names)
            }
        except Exception as e:
            logger.warning(f"提取Excel文档结构失败: {str(e)}")
            return {'headings': [], 'sections': []}
    
    def _extract_markdown_structure(self, file_path: str) -> Dict[str, Any]:
        """提取Markdown文档结构"""
        try:
            content = self._read_markdown_file(file_path)
            lines = content.split('\n')
            
            headings = []
            sections = []
            
            for line in lines:
                line = line.strip()
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    text = line.lstrip('#').strip()
                    headings.append({
                        'level': level,
                        'text': text
                    })
                elif line:
                    sections.append({
                        'type': 'paragraph',
                        'text': line
                    })
            
            return {
                'headings': headings,
                'sections': sections
            }
        except Exception as e:
            logger.warning(f"提取Markdown文档结构失败: {str(e)}")
            return {'headings': [], 'sections': []}
